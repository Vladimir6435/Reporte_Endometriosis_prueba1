import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Configuración de la página
st.set_page_config(
    page_title="Reporte Ultrasonido Endometriosis #Enzian Asociación Costarricense de Ginecología",
    page_icon="🏥",
    layout="wide"
)

# CSS personalizado
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        color: #2c3e50;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .section-header {
        background-color: #f0f2f6;
        padding: 0.5rem;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Título principal
st.markdown('<h1 class="main-header">📊 Sistema de Reporte Ultrasonográfico de Endometriosis<br>Clasificación #Enzian</h1>', unsafe_allow_html=True)

# Inicializar session state
if 'data' not in st.session_state:
    st.session_state.data = {
        'paciente': {},
        'peritoneo': {},
        'ovarios': {'izquierdo': {}, 'derecho': {}},
        'tubos': {'izquierdo': {}, 'derecho': {}},
        'compartimento_a': {},
        'compartimento_b': {'izquierdo': {}, 'derecho': {}},
        'compartimento_c': {},
        'localizaciones_f': {}
    }

# Funciones de validación
def calcular_clasificacion_ovario(diametro):
    """Calcula la clasificación O según el diámetro"""
    if diametro < 3:
        return "O1"
    elif 3 <= diametro <= 7:
        return "O2"
    else:
        return "O3"

def calcular_clasificacion_compartimento(medida):
    """Calcula clasificación para compartimentos A, B, C"""
    if medida < 1:
        return "1"
    elif 1 <= medida <= 3:
        return "2"
    else:
        return "3"

def validar_consistencia(compartimento, medida, clasificacion_manual):
    """Valida que la clasificación manual coincida con la medida"""
    if compartimento in ['A', 'B', 'C']:
        clasificacion_calculada = calcular_clasificacion_compartimento(medida)
        if clasificacion_calculada != clasificacion_manual:
            return False, f"⚠️ Inconsistencia: La medida {medida}cm sugiere clasificación {clasificacion_calculada}, pero seleccionaste {clasificacion_manual}"
    return True, ""

# Pestañas principales
tabs = st.tabs([
    "👤 Datos del Paciente",
    "🔴 Peritoneo (P)",
    "🥚 Ovarios (O)",
    "🎗️ Condición Tubo-Ovárica (T)",
    "🅰️ Compartimento A",
    "🅱️ Compartimento B",
    "🅲 Compartimento C",
    "📍 Localizaciones F",
    "📋 Generar Reporte"
])

# ============= PESTAÑA 1: DATOS DEL PACIENTE =============
with tabs[0]:
    st.markdown('<div class="section-header"><h2>Datos del Paciente</h2></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        nombre = st.text_input("Nombre completo", key="nombre_paciente")
        edad = st.number_input("Edad", min_value=0, max_value=120, key="edad_paciente")
        cedula = st.text_input("Número de identificación", key="cedula_paciente")
        
    with col2:
        fecha_estudio = st.date_input("Fecha del estudio", datetime.now())
        medico = st.text_input("Médico solicitante", key="medico_solicitante")
        indicacion = st.text_area("Indicación del estudio", key="indicacion_estudio")
    
    st.session_state.data['paciente'] = {
        'nombre': nombre,
        'edad': edad,
        'cedula': cedula,
        'fecha': fecha_estudio,
        'medico': medico,
        'indicacion': indicacion
    }

# ============= PESTAÑA 2: PERITONEO (P) =============
with tabs[1]:
    st.markdown('<div class="section-header"><h2>🔴 Peritoneo (P)</h2></div>', unsafe_allow_html=True)
    st.info("📌 Lesiones superficiales peritoneales (<5mm de invasión subperitoneal)")
    
    peritoneo_estado = st.radio(
        "Estado del peritoneo:",
        ["Normal", "Anormal"],
        key="peritoneo_estado",
        horizontal=True
    )
    
    if peritoneo_estado == "Anormal":
        col1, col2 = st.columns(2)
        
        with col1:
            clasificacion_p = st.select_slider(
                "Clasificación según diámetro virtual (suma de lesiones):",
                options=["P1 (<3 cm)", "P2 (3-7 cm)", "P3 (>7 cm)"],
                key="clasificacion_p"
            )
            
        with col2:
            diametro_total = st.number_input(
                "Diámetro total aproximado (cm):",
                min_value=0.0,
                max_value=20.0,
                step=0.1,
                key="diametro_peritoneo"
            )
        
        localizaciones = st.multiselect(
            "Localizaciones afectadas:",
            ["Fondo de saco de Douglas", "Peritoneo pélvico lateral", 
             "Ligamento ancho", "Peritoneo vesical", "Otras"],
            key="localizaciones_peritoneo"
        )
        
        descripcion_p = st.text_area(
            "Descripción adicional:",
            key="descripcion_peritoneo"
        )
        
        st.session_state.data['peritoneo'] = {
            'estado': 'anormal',
            'clasificacion': clasificacion_p,
            'diametro': diametro_total,
            'localizaciones': localizaciones,
            'descripcion': descripcion_p
        }
    else:
        st.session_state.data['peritoneo'] = {'estado': 'normal'}

# ============= PESTAÑA 3: OVARIOS (O) =============
with tabs[2]:
    st.markdown('<div class="section-header"><h2>🥚 Ovarios (O)</h2></div>', unsafe_allow_html=True)
    st.info("📌 Incluye endometriomas y focos infiltrantes de superficie ovárica (≥5mm)")
    
    # Ovario derecho
    st.markdown("### Ovario Derecho")
    ovario_der_estado = st.radio(
        "Estado:",
        ["Normal", "Anormal", "No visualizado"],
        key="ovario_der_estado",
        horizontal=True
    )
    
    if ovario_der_estado == "Anormal":
        col1, col2, col3 = st.columns(3)
        
        with col1:
            diametro_der = st.number_input(
                "Diámetro máximo (cm):",
                min_value=0.0,
                max_value=15.0,
                step=0.1,
                key="diametro_ovario_der"
            )
            
            # Calcular clasificación automática
            if diametro_der > 0:
                clasificacion_auto = calcular_clasificacion_ovario(diametro_der)
                st.info(f"💡 Clasificación sugerida: {clasificacion_auto}")
            
        with col2:
            num_endometriomas_der = st.number_input(
                "Número de endometriomas:",
                min_value=1,
                max_value=10,
                step=1,
                key="num_endometriomas_der"
            )
            
        with col3:
            clasificacion_o_der = st.select_slider(
                "Clasificación:",
                options=["O1 (<3cm)", "O2 (3-7cm)", "O3 (>7cm)"],
                key="clasificacion_o_der"
            )
        
        # Criterios IOTA
        st.markdown("#### Criterios IOTA")
        col1, col2 = st.columns(2)
        
        with col1:
            estructura_der = st.selectbox(
                "Estructura:",
                ["Unilocular", "Multilocular", "Unilocular-sólido", "Multilocular-sólido", "Sólido"],
                key="estructura_ovario_der"
            )
            
            contenido_der = st.selectbox(
                "Contenido:",
                ["Anecoico", "Homogéneo de baja intensidad (ground glass)", 
                 "Heterogéneo", "Con nivel líquido-líquido"],
                key="contenido_ovario_der"
            )
            
        with col2:
            vascularizacion_der = st.selectbox(
                "Vascularización al Doppler:",
                ["Ausente", "Mínima periférica", "Moderada", "Abundante"],
                key="vascularizacion_der"
            )
            
            adherencias_der = st.checkbox(
                "Signos de adherencias a estructuras adyacentes",
                key="adherencias_ovario_der"
            )
        
        descripcion_ovario_der = st.text_area(
            "Descripción adicional del ovario derecho:",
            key="descripcion_ovario_der"
        )
        
        st.session_state.data['ovarios']['derecho'] = {
            'estado': 'anormal',
            'diametro': diametro_der,
            'num_endometriomas': num_endometriomas_der,
            'clasificacion': clasificacion_o_der,
            'estructura': estructura_der,
            'contenido': contenido_der,
            'vascularizacion': vascularizacion_der,
            'adherencias': adherencias_der,
            'descripcion': descripcion_ovario_der
        }
    elif ovario_der_estado == "No visualizado":
        st.session_state.data['ovarios']['derecho'] = {'estado': 'no_visualizado'}
    else:
        st.session_state.data['ovarios']['derecho'] = {'estado': 'normal'}
    
    st.markdown("---")
    
    # Ovario izquierdo
    st.markdown("### Ovario Izquierdo")
    ovario_izq_estado = st.radio(
        "Estado:",
        ["Normal", "Anormal", "No visualizado"],
        key="ovario_izq_estado",
        horizontal=True
    )
    
    if ovario_izq_estado == "Anormal":
        col1, col2, col3 = st.columns(3)
        
        with col1:
            diametro_izq = st.number_input(
                "Diámetro máximo (cm):",
                min_value=0.0,
                max_value=15.0,
                step=0.1,
                key="diametro_ovario_izq"
            )
            
            if diametro_izq > 0:
                clasificacion_auto = calcular_clasificacion_ovario(diametro_izq)
                st.info(f"💡 Clasificación sugerida: {clasificacion_auto}")
            
        with col2:
            num_endometriomas_izq = st.number_input(
                "Número de endometriomas:",
                min_value=1,
                max_value=10,
                step=1,
                key="num_endometriomas_izq"
            )
            
        with col3:
            clasificacion_o_izq = st.select_slider(
                "Clasificación:",
                options=["O1 (<3cm)", "O2 (3-7cm)", "O3 (>7cm)"],
                key="clasificacion_o_izq"
            )
        
        st.markdown("#### Criterios IOTA")
        col1, col2 = st.columns(2)
        
        with col1:
            estructura_izq = st.selectbox(
                "Estructura:",
                ["Unilocular", "Multilocular", "Unilocular-sólido", "Multilocular-sólido", "Sólido"],
                key="estructura_ovario_izq"
            )
            
            contenido_izq = st.selectbox(
                "Contenido:",
                ["Anecoico", "Homogéneo de baja intensidad (ground glass)", 
                 "Heterogéneo", "Con nivel líquido-líquido"],
                key="contenido_ovario_izq"
            )
            
        with col2:
            vascularizacion_izq = st.selectbox(
                "Vascularización al Doppler:",
                ["Ausente", "Mínima periférica", "Moderada", "Abundante"],
                key="vascularizacion_izq"
            )
            
            adherencias_izq = st.checkbox(
                "Signos de adherencias a estructuras adyacentes",
                key="adherencias_ovario_izq"
            )
        
        descripcion_ovario_izq = st.text_area(
            "Descripción adicional del ovario izquierdo:",
            key="descripcion_ovario_izq"
        )
        
        st.session_state.data['ovarios']['izquierdo'] = {
            'estado': 'anormal',
            'diametro': diametro_izq,
            'num_endometriomas': num_endometriomas_izq,
            'clasificacion': clasificacion_o_izq,
            'estructura': estructura_izq,
            'contenido': contenido_izq,
            'vascularizacion': vascularizacion_izq,
            'adherencias': adherencias_izq,
            'descripcion': descripcion_ovario_izq
        }
    elif ovario_izq_estado == "No visualizado":
        st.session_state.data['ovarios']['izquierdo'] = {'estado': 'no_visualizado'}
    else:
        st.session_state.data['ovarios']['izquierdo'] = {'estado': 'normal'}

# ============= PESTAÑA 4: CONDICIÓN TUBO-OVÁRICA (T) =============
with tabs[3]:
    st.markdown('<div class="section-header"><h2>🎗️ Condición Tubo-Ovárica (T)</h2></div>', unsafe_allow_html=True)
    st.info("📌 Evaluación de adherencias y movilidad tubo-ovárica mediante sliding sign")
    
    # Lado derecho
    st.markdown("### Lado Derecho")
    tubo_der_estado = st.radio(
        "Estado:",
        ["Normal - Movilidad preservada", "Anormal - Adherencias presentes", "No evaluable"],
        key="tubo_der_estado",
        horizontal=True
    )
    
    if tubo_der_estado == "Anormal - Adherencias presentes":
        clasificacion_t_der = st.select_slider(
            "Clasificación:",
            options=[
                "T1 - Adherencias ovario-pared pélvica",
                "T2 - T1 + adherencias al útero",
                "T3 - T2 + adherencias a LSU/intestino"
            ],
            key="clasificacion_t_der"
        )
        
        sliding_sign_der = st.select_slider(
            "Sliding sign:",
            options=["Positivo (móvil)", "Limitado", "Negativo (fijo)"],
            key="sliding_sign_der"
        )
        
        permeabilidad_der = st.radio(
            "Permeabilidad tubárica (opcional):",
            ["No evaluada", "Permeable (+)", "No permeable (-)"],
            key="permeabilidad_der",
            horizontal=True
        )
        
        descripcion_tubo_der = st.text_area(
            "Descripción adicional lado derecho:",
            key="descripcion_tubo_der"
        )
        
        st.session_state.data['tubos']['derecho'] = {
            'estado': 'anormal',
            'clasificacion': clasificacion_t_der,
            'sliding_sign': sliding_sign_der,
            'permeabilidad': permeabilidad_der,
            'descripcion': descripcion_tubo_der
        }
    elif tubo_der_estado == "No evaluable":
        st.session_state.data['tubos']['derecho'] = {'estado': 'no_evaluable'}
    else:
        st.session_state.data['tubos']['derecho'] = {'estado': 'normal'}
    
    st.markdown("---")
    
    # Lado izquierdo
    st.markdown("### Lado Izquierdo")
    tubo_izq_estado = st.radio(
        "Estado:",
        ["Normal - Movilidad preservada", "Anormal - Adherencias presentes", "No evaluable"],
        key="tubo_izq_estado",
        horizontal=True
    )
    
    if tubo_izq_estado == "Anormal - Adherencias presentes":
        clasificacion_t_izq = st.select_slider(
            "Clasificación:",
            options=[
                "T1 - Adherencias ovario-pared pélvica",
                "T2 - T1 + adherencias al útero",
                "T3 - T2 + adherencias a LSU/intestino"
            ],
            key="clasificacion_t_izq"
        )
        
        sliding_sign_izq = st.select_slider(
            "Sliding sign:",
            options=["Positivo (móvil)", "Limitado", "Negativo (fijo)"],
            key="sliding_sign_izq"
        )
        
        permeabilidad_izq = st.radio(
            "Permeabilidad tubárica (opcional):",
            ["No evaluada", "Permeable (+)", "No permeable (-)"],
            key="permeabilidad_izq",
            horizontal=True
        )
        
        descripcion_tubo_izq = st.text_area(
            "Descripción adicional lado izquierdo:",
            key="descripcion_tubo_izq"
        )
        
        st.session_state.data['tubos']['izquierdo'] = {
            'estado': 'anormal',
            'clasificacion': clasificacion_t_izq,
            'sliding_sign': sliding_sign_izq,
            'permeabilidad': permeabilidad_izq,
            'descripcion': descripcion_tubo_izq
        }
    elif tubo_izq_estado == "No evaluable":
        st.session_state.data['tubos']['izquierdo'] = {'estado': 'no_evaluable'}
    else:
        st.session_state.data['tubos']['izquierdo'] = {'estado': 'normal'}

# ============= PESTAÑA 5: COMPARTIMENTO A =============
with tabs[4]:
    st.markdown('<div class="section-header"><h2>🅰️ Compartimento A</h2></div>', unsafe_allow_html=True)
    st.info("📌 Vagina, espacio rectovaginal y área retrocervical (eje craneocaudal)")
    
    comp_a_estado = st.radio(
        "Estado:",
        ["Normal", "Anormal"],
        key="comp_a_estado",
        horizontal=True
    )
    
    if comp_a_estado == "Anormal":
        col1, col2 = st.columns(2)
        
        with col1:
            diametro_a = st.number_input(
                "Diámetro máximo en plano sagital medio (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="diametro_comp_a"
            )
            
            if diametro_a > 0:
                clasificacion_sugerida = calcular_clasificacion_compartimento(diametro_a)
                st.info(f"💡 Clasificación sugerida: A{clasificacion_sugerida}")
            
        with col2:
            clasificacion_a = st.select_slider(
                "Clasificación:",
                options=["A1 (<1 cm)", "A2 (1-3 cm)", "A3 (>3 cm)"],
                key="clasificacion_a"
            )
        
        # Validación de consistencia
        if diametro_a > 0:
            clase_manual = clasificacion_a[1]  # Extraer el número
            es_valido, mensaje = validar_consistencia('A', diametro_a, clase_manual)
            if not es_valido:
                st.warning(mensaje)
        
        localizacion_a = st.multiselect(
            "Localización específica:",
            ["Fórnix vaginal posterior", "Espacio rectovaginal", "Área retrocervical"],
            key="localizacion_comp_a"
        )
        
        ecogenicidad_a = st.selectbox(
            "Ecogenicidad de la lesión:",
            ["Hipoecogénica", "Isoecogénica", "Heterogénea"],
            key="ecogenicidad_comp_a"
        )
        
        contornos_a = st.selectbox(
            "Contornos:",
            ["Regulares", "Irregulares", "Espiculados"],
            key="contornos_comp_a"
        )
        
        descripcion_a = st.text_area(
            "Descripción adicional:",
            key="descripcion_comp_a"
        )
        
        st.session_state.data['compartimento_a'] = {
            'estado': 'anormal',
            'diametro': diametro_a,
            'clasificacion': clasificacion_a,
            'localizacion': localizacion_a,
            'ecogenicidad': ecogenicidad_a,
            'contornos': contornos_a,
            'descripcion': descripcion_a
        }
    else:
        st.session_state.data['compartimento_a'] = {'estado': 'normal'}

# ============= PESTAÑA 6: COMPARTIMENTO B =============
with tabs[5]:
    st.markdown('<div class="section-header"><h2>🅱️ Compartimento B</h2></div>', unsafe_allow_html=True)
    st.info("📌 Ligamentos uterosacros, ligamentos cardinales y pared pélvica lateral (eje mediolateral)")
    
    # Lado derecho
    st.markdown("### Ligamento Uterosacro Derecho")
    lsu_der_estado = st.radio(
        "Estado:",
        ["Normal", "Anormal"],
        key="lsu_der_estado",
        horizontal=True
    )
    
    if lsu_der_estado == "Anormal":
        col1, col2, col3 = st.columns(3)
        
        with col1:
            diametro_max_b_der = st.number_input(
                "Diámetro máximo (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="diametro_lsu_der"
            )
            
            if diametro_max_b_der > 0:
                clasificacion_sugerida = calcular_clasificacion_compartimento(diametro_max_b_der)
                st.info(f"💡 Clasificación sugerida: B{clasificacion_sugerida}")
        
        with col2:
            dim_ap_b_der = st.number_input(
                "Dimensión anteroposterior (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="dim_ap_lsu_der"
            )
        
        with col3:
            dim_cc_b_der = st.number_input(
                "Dimensión craneocaudal (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="dim_cc_lsu_der"
            )
        
        clasificacion_b_der = st.select_slider(
            "Clasificación:",
            options=["B1 (<1 cm)", "B2 (1-3 cm)", "B3 (>3 cm)"],
            key="clasificacion_b_der"
        )
        
        # Validación
        if diametro_max_b_der > 0:
            clase_manual = clasificacion_b_der[1]
            es_valido, mensaje = validar_consistencia('B', diametro_max_b_der, clase_manual)
            if not es_valido:
                st.warning(mensaje)
        
        st.markdown("#### Evaluación de movilidad (Sliding Sign)")
        sliding_lsu_der = st.select_slider(
            "Sliding sign del LSU derecho:",
            options=["Positivo (móvil)", "Limitado", "Negativo (fijo)"],
            key="sliding_lsu_der"
        )
        
        distancia_cervix_der = st.number_input(
            "Distancia desde inserción cervical (cm):",
            min_value=0.0,
            max_value=10.0,
            step=0.1,
            key="distancia_cervix_der"
        )
        
        descripcion_lsu_der = st.text_area(
            "Descripción adicional LSU derecho:",
            key="descripcion_lsu_der"
        )
        
        st.session_state.data['compartimento_b']['derecho'] = {
            'estado': 'anormal',
            'diametro_max': diametro_max_b_der,
            'dim_ap': dim_ap_b_der,
            'dim_cc': dim_cc_b_der,
            'clasificacion': clasificacion_b_der,
            'sliding_sign': sliding_lsu_der,
            'distancia_cervix': distancia_cervix_der,
            'descripcion': descripcion_lsu_der
        }
    else:
        st.session_state.data['compartimento_b']['derecho'] = {'estado': 'normal'}
    
    st.markdown("---")
    
    # Lado izquierdo
    st.markdown("### Ligamento Uterosacro Izquierdo")
    lsu_izq_estado = st.radio(
        "Estado:",
        ["Normal", "Anormal"],
        key="lsu_izq_estado",
        horizontal=True
    )
    
    if lsu_izq_estado == "Anormal":
        col1, col2, col3 = st.columns(3)
        
        with col1:
            diametro_max_b_izq = st.number_input(
                "Diámetro máximo (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="diametro_lsu_izq"
            )
            
            if diametro_max_b_izq > 0:
                clasificacion_sugerida = calcular_clasificacion_compartimento(diametro_max_b_izq)
                st.info(f"💡 Clasificación sugerida: B{clasificacion_sugerida}")
        
        with col2:
            dim_ap_b_izq = st.number_input(
                "Dimensión anteroposterior (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="dim_ap_lsu_izq"
            )
        
        with col3:
            dim_cc_b_izq = st.number_input(
                "Dimensión craneocaudal (cm):",
                min_value=0.0,
                max_value=10.0,
                step=0.1,
                key="dim_cc_lsu_izq"
            )
        
        clasificacion_b_izq = st.select_slider(
            "Clasificación:",
            options=["B1 (<1 cm)", "B2 (1-3 cm)", "B3 (>3 cm)"],
            key="clasificacion_b_izq"
        )
        
        # Validación
        if diametro_max_b_izq > 0:
            clase_manual = clasificacion_b_izq[1]
            es_valido, mensaje = validar_consistencia('B', diametro_max_b_izq, clase_manual)
            if not es_valido:
                st.warning(mensaje)
        
        st.markdown("#### Evaluación de movilidad (Sliding Sign)")
        sliding_lsu_izq = st.select_slider(
            "Sliding sign del LSU izquierdo:",
            options=["Positivo (móvil)", "Limitado", "Negativo (fijo)"],
            key="sliding_lsu_izq"
        )
        
        distancia_cervix_izq = st.number_input(
            "Distancia desde inserción cervical (cm):",
            min_value=0.0,
            max_value=10.0,
            step=0.1,
            key="distancia_cervix_izq"
        )
        
        descripcion_lsu_izq = st.text_area(
            "Descripción adicional LSU izquierdo:",
            key="descripcion_lsu_izq"
        )
        
        st.session_state.data['compartimento_b']['izquierdo'] = {
            'estado': 'anormal',
            'diametro_max': diametro_max_b_izq,
            'dim_ap': dim_ap_b_izq,
            'dim_cc': dim_cc_b_izq,
            'clasificacion': clasificacion_b_izq,
            'sliding_sign': sliding_lsu_izq,
            'distancia_cervix': distancia_cervix_izq,
            'descripcion': descripcion_lsu_izq
        }
    else:
        st.session_state.data['compartimento_b']['izquierdo'] = {'estado': 'normal'}

# ============= PESTAÑA 7: COMPARTIMENTO C =============
with tabs[6]:
    st.markdown('<div class="section-header"><h2>🅲 Compartimento C</h2></div>', unsafe_allow_html=True)
    st.info("📌 Recto (hasta 16 cm del margen anal) - Eje ventrodorsal")
    
    comp_c_estado = st.radio(
        "Estado:",
        ["Normal", "Anormal"],
        key="comp_c_estado",
        horizontal=True
    )
    
    if comp_c_estado == "Anormal":
        col1, col2 = st.columns(2)
        
        with col1:
            longitud_lesion_c = st.number_input(
                "Longitud de la lesión (cm):",
                min_value=0.0,
                max_value=20.0,
                step=0.1,
                key="longitud_lesion_c"
            )
            
            if longitud_lesion_c > 0:
                clasificacion_sugerida = calcular_clasificacion_compartimento(longitud_lesion_c)
                st.info(f"💡 Clasificación sugerida: C{clasificacion_sugerida}")
        
        with col2:
            clasificacion_c = st.select_slider(
                "Clasificación:",
                options=["C1 (<1 cm)", "C2 (1-3 cm)", "C3 (>3 cm)"],
                key="clasificacion_c"
            )
        
        # Validación
        if longitud_lesion_c > 0:
            clase_manual = clasificacion_c[1]
            es_valido, mensaje = validar_consistencia('C', longitud_lesion_c, clase_manual)
            if not es_valido:
                st.warning(mensaje)
        
        st.markdown("#### Características específicas del recto")
        
        distancia_anal = st.number_input(
            "Distancia desde margen anal (cm):",
            min_value=0.0,
            max_value=16.0,
            step=0.5,
            key="distancia_anal_c",
            help="Hasta 16 cm = recto; >16 cm = clasificar como FI"
        )
        
        if distancia_anal > 16:
            st.error("⚠️ Lesiones >16 cm del margen anal deben clasificarse como FI (sigma)")
        
        profundidad_infiltracion = st.selectbox(
            "Profundidad de infiltración:",
            ["Serosa/subserosa", "Muscular propia", "Submucosa", "Mucosa"],
            key="profundidad_infiltracion_c"
        )
        
        porcentaje_circunferencia = st.slider(
            "Porcentaje de circunferencia afectada:",
            min_value=0,
            max_value=100,
            step=5,
            value=0,
            key="porcentaje_circunferencia_c"
        )
        
        st.write(f"Circunferencia afectada: {porcentaje_circunferencia}%")
        
        estenosis = st.checkbox(
            "Signos de estenosis",
            key="estenosis_c"
        )
        
        sliding_sign_rectal = st.select_slider(
            "Sliding sign rectal:",
            options=["Positivo (móvil)", "Limitado", "Negativo (fijo)"],
            key="sliding_sign_rectal"
        )
        
        descripcion_c = st.text_area(
            "Descripción adicional:",
            key="descripcion_comp_c"
        )
        
        st.session_state.data['compartimento_c'] = {
            'estado': 'anormal',
            'longitud': longitud_lesion_c,
            'clasificacion': clasificacion_c,
            'distancia_anal': distancia_anal,
            'profundidad': profundidad_infiltracion,
            'circunferencia': porcentaje_circunferencia,
            'estenosis': estenosis,
            'sliding_sign': sliding_sign_rectal,
            'descripcion': descripcion_c
        }
    else:
        st.session_state.data['compartimento_c'] = {'estado': 'normal'}

# ============= PESTAÑA 8: LOCALIZACIONES F =============
with tabs[7]:
    st.markdown('<div class="section-header"><h2>📍 Localizaciones F (Far locations)</h2></div>', unsafe_allow_html=True)
    st.info("📌 Localizaciones extragenitales y otras ubicaciones distantes")
    
    # Adenomiosis (FA)
    st.markdown("### FA - Adenomiosis")
    adenomiosis = st.radio(
        "¿Adenomiosis presente?",
        ["No", "Sí"],
        key="adenomiosis_presente",
        horizontal=True
    )
    
    if adenomiosis == "Sí":
        criterios_musa = st.multiselect(
            "Criterios MUSA presentes:",
            [
                "Asimetría de paredes miometriales",
                "Quistes miometriales",
                "Hiperplasia endometrial focal",
                "Líneas de sombra",
                "Áreas heterogéneas en miometrio",
                "Zona juncional irregular",
                "Vascularización translesional"
            ],
            key="criterios_musa"
        )
        
        descripcion_fa = st.text_area(
            "Descripción de adenomiosis:",
            key="descripcion_adenomiosis"
        )
        
        st.session_state.data['localizaciones_f']['adenomiosis'] = {
            'presente': True,
            'criterios_musa': criterios_musa,
            'descripcion': descripcion_fa
        }
    else:
        st.session_state.data['localizaciones_f']['adenomiosis'] = {'presente': False}
    
    st.markdown("---")
    
    # Vejiga (FB)
    st.markdown("### FB - Vejiga")
    vejiga = st.radio(
        "¿Compromiso vesical?",
        ["No", "Sí"],
        key="vejiga_presente",
        horizontal=True
    )
    
    if vejiga == "Sí":
        localizacion_vejiga = st.selectbox(
            "Localización en vejiga:",
            ["Pared posterior", "Cúpula", "Trígono", "Otras"],
            key="localizacion_vejiga"
        )
        
        profundidad_vejiga = st.selectbox(
            "Profundidad:",
            ["Serosa", "Muscular", "Submucosa", "Mucosa"],
            key="profundidad_vejiga"
        )
        
        dimension_vejiga = st.number_input(
            "Dimensión máxima (cm):",
            min_value=0.0,
            max_value=10.0,
            step=0.1,
            key="dimension_vejiga"
        )
        
        descripcion_fb = st.text_area(
            "Descripción de compromiso vesical:",
            key="descripcion_vejiga"
        )
        
        st.session_state.data['localizaciones_f']['vejiga'] = {
            'presente': True,
            'localizacion': localizacion_vejiga,
            'profundidad': profundidad_vejiga,
            'dimension': dimension_vejiga,
            'descripcion': descripcion_fb
        }
    else:
        st.session_state.data['localizaciones_f']['vejiga'] = {'presente': False}
    
    st.markdown("---")
    
    # Uréter (FU)
    st.markdown("### FU - Uréter")
    ureter = st.radio(
        "¿Compromiso ureteral?",
        ["No", "Sí"],
        key="ureter_presente",
        horizontal=True
    )
    
    if ureter == "Sí":
        lado_ureter = st.multiselect(
            "Lado(s) afectado(s):",
            ["Derecho", "Izquierdo"],
            key="lado_ureter"
        )
        
        for lado in lado_ureter:
            st.markdown(f"#### Uréter {lado}")
            col1, col2 = st.columns(2)
            
            with col1:
                diametro_ureter = st.number_input(
                    f"Diámetro uréter {lado.lower()} (mm):",
                    min_value=0.0,
                    max_value=20.0,
                    step=0.5,
                    key=f"diametro_ureter_{lado.lower()}"
                )
                
                if diametro_ureter >= 6:
                    st.warning(f"⚠️ Dilatación ureteral (≥6mm) en lado {lado.lower()}")
            
            with col2:
                hidronefrosis = st.selectbox(
                    f"Hidronefrosis {lado.lower()}:",
                    ["Ausente", "Leve", "Moderada", "Severa"],
                    key=f"hidronefrosis_{lado.lower()}"
                )
        
        tipo_compromiso = st.selectbox(
            "Tipo de compromiso:",
            ["Extrínseco", "Intrínseco", "Mixto"],
            key="tipo_compromiso_ureter"
        )
        
        descripcion_fu = st.text_area(
            "Descripción de compromiso ureteral:",
            key="descripcion_ureter"
        )
        
        st.session_state.data['localizaciones_f']['ureter'] = {
            'presente': True,
            'lados': lado_ureter,
            'tipo_compromiso': tipo_compromiso,
            'descripcion': descripcion_fu
        }
    else:
        st.session_state.data['localizaciones_f']['ureter'] = {'presente': False}
    
    st.markdown("---")
    
    # Intestino (FI)
    st.markdown("### FI - Intestino (>16 cm del margen anal)")
    intestino = st.radio(
        "¿Compromiso intestinal extra-rectal?",
        ["No", "Sí"],
        key="intestino_presente",
        horizontal=True
    )
    
    if intestino == "Sí":
        localizacion_intestino = st.multiselect(
            "Localización(es):",
            ["Sigma (>16cm)", "Colon transverso", "Ciego", "Apéndice", "Intestino delgado"],
            key="localizacion_intestino"
        )
        
        dimension_intestino = st.number_input(
            "Dimensión máxima de la lesión (cm):",
            min_value=0.0,
            max_value=15.0,
            step=0.1,
            key="dimension_intestino"
        )
        
        descripcion_fi = st.text_area(
            "Descripción de compromiso intestinal:",
            key="descripcion_intestino"
        )
        
        st.session_state.data['localizaciones_f']['intestino'] = {
            'presente': True,
            'localizaciones': localizacion_intestino,
            'dimension': dimension_intestino,
            'descripcion': descripcion_fi
        }
    else:
        st.session_state.data['localizaciones_f']['intestino'] = {'presente': False}
    
    st.markdown("---")
    
    # Otras localizaciones (FO)
    st.markdown("### F(...) - Otras Localizaciones")
    otras_localizaciones = st.radio(
        "¿Otras localizaciones?",
        ["No", "Sí"],
        key="otras_localizaciones_presente",
        horizontal=True
    )
    
    if otras_localizaciones == "Sí":
        tipos_otras = st.multiselect(
            "Seleccione localización(es):",
            ["Pared abdominal", "Diafragma", "Pulmón", "Nervio", "Cicatriz quirúrgica", "Ombligo", "Otras"],
            key="tipos_otras_localizaciones"
        )
        
        for tipo in tipos_otras:
            descripcion_otra = st.text_area(
                f"Descripción de {tipo}:",
                key=f"descripcion_otra_{tipo.replace(' ', '_')}"
            )
        
        st.session_state.data['localizaciones_f']['otras'] = {
            'presente': True,
            'tipos': tipos_otras
        }
    else:
        st.session_state.data['localizaciones_f']['otras'] = {'presente': False}

# ============= PESTAÑA 9: GENERAR REPORTE =============
with tabs[8]:
    st.markdown('<div class="section-header"><h2>📋 Generar Reporte Final</h2></div>', unsafe_allow_html=True)
    
    # Función para generar código #Enzian
    def generar_codigo_enzian():
        codigo = "#Enzian(u) "
        componentes = []
        
        # Peritoneo (P)
        if st.session_state.data['peritoneo'].get('estado') == 'anormal':
            clasificacion = st.session_state.data['peritoneo'].get('clasificacion', 'P1')
            componentes.append(clasificacion.split()[0])  # Extrae "P1", "P2", o "P3"
        
        # Ovarios (O)
        ovario_izq = st.session_state.data['ovarios']['izquierdo']
        ovario_der = st.session_state.data['ovarios']['derecho']
        
        if ovario_izq.get('estado') == 'anormal' or ovario_der.get('estado') == 'anormal':
            clase_izq = "0"
            clase_der = "0"
            
            if ovario_izq.get('estado') == 'anormal':
                clase_izq = ovario_izq.get('clasificacion', 'O1')[1]  # Extrae el número
            elif ovario_izq.get('estado') == 'no_visualizado':
                clase_izq = "x"
                
            if ovario_der.get('estado') == 'anormal':
                clase_der = ovario_der.get('clasificacion', 'O1')[1]
            elif ovario_der.get('estado') == 'no_visualizado':
                clase_der = "x"
                
            componentes.append(f"O{clase_izq}/{clase_der}")
        
        # Tubos (T)
        tubo_izq = st.session_state.data['tubos']['izquierdo']
        tubo_der = st.session_state.data['tubos']['derecho']
        
        if tubo_izq.get('estado') == 'anormal' or tubo_der.get('estado') == 'anormal':
            clase_izq = "0"
            clase_der = "0"
            
            if tubo_izq.get('estado') == 'anormal':
                clase_texto = tubo_izq.get('clasificacion', 'T1')
                clase_izq = clase_texto[1]  # Extrae el número
                
            if tubo_der.get('estado') == 'anormal':
                clase_texto = tubo_der.get('clasificacion', 'T1')
                clase_der = clase_texto[1]
                
            componentes.append(f"T{clase_izq}/{clase_der}")
        
        # Compartimento A
        if st.session_state.data['compartimento_a'].get('estado') == 'anormal':
            clase_a = st.session_state.data['compartimento_a'].get('clasificacion', 'A1')
            componentes.append(clase_a.split()[0])
        
        # Compartimento B
        lsu_izq = st.session_state.data['compartimento_b']['izquierdo']
        lsu_der = st.session_state.data['compartimento_b']['derecho']
        
        if lsu_izq.get('estado') == 'anormal' or lsu_der.get('estado') == 'anormal':
            clase_izq = "0"
            clase_der = "0"
            
            if lsu_izq.get('estado') == 'anormal':
                clase_izq = lsu_izq.get('clasificacion', 'B1')[1]
                
            if lsu_der.get('estado') == 'anormal':
                clase_der = lsu_der.get('clasificacion', 'B1')[1]
                
            componentes.append(f"B{clase_izq}/{clase_der}")
        
        # Compartimento C
        if st.session_state.data['compartimento_c'].get('estado') == 'anormal':
            clase_c = st.session_state.data['compartimento_c'].get('clasificacion', 'C1')
            componentes.append(clase_c.split()[0])
        
        # Localizaciones F
        loc_f = st.session_state.data['localizaciones_f']
        
        if loc_f.get('adenomiosis', {}).get('presente'):
            componentes.append("FA")
            
        if loc_f.get('vejiga', {}).get('presente'):
            componentes.append("FB")
            
        if loc_f.get('ureter', {}).get('presente'):
            lados = loc_f['ureter'].get('lados', [])
            for lado in lados:
                inicial = 'r' if lado == 'Derecho' else 'l'
                componentes.append(f"FU({inicial})")
                
        if loc_f.get('intestino', {}).get('presente'):
            locs = loc_f['intestino'].get('localizaciones', [])
            for loc in locs:
                if 'Sigma' in loc:
                    componentes.append("FI(Sigma)")
                elif 'Apéndice' in loc:
                    componentes.append("FI(Apéndice)")
                else:
                    componentes.append(f"FI({loc})")
                    
        if loc_f.get('otras', {}).get('presente'):
            tipos = loc_f['otras'].get('tipos', [])
            for tipo in tipos:
                componentes.append(f"F({tipo})")
        
        codigo += ", ".join(componentes) if componentes else "Sin hallazgos de endometriosis"
        
        return codigo
    
    # Función para generar reporte en Word
    def generar_reporte_word():
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Encabezado
        header = doc.add_heading('REPORTE ULTRASONOGRÁFICO', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subheader = doc.add_heading('Evaluación de Endometriosis Asociación Costarricense de Ginecologia - Clasificación #Enzian', level=2)
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Datos del paciente
        doc.add_heading('DATOS DEL PACIENTE', level=1)
        paciente = st.session_state.data['paciente']
        
        tabla_paciente = doc.add_table(rows=5, cols=2)
        tabla_paciente.style = 'Light Grid Accent 1'
        
        datos = [
            ('Nombre:', paciente.get('nombre', 'N/A')),
            ('Identificación:', paciente.get('cedula', 'N/A')),
            ('Edad:', f"{paciente.get('edad', 'N/A')} años"),
            ('Fecha del estudio:', str(paciente.get('fecha', 'N/A'))),
            ('Médico solicitante:', paciente.get('medico', 'N/A'))
        ]
        
        for i, (campo, valor) in enumerate(datos):
            tabla_paciente.rows[i].cells[0].text = campo
            tabla_paciente.rows[i].cells[1].text = str(valor)
        
        if paciente.get('indicacion'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Indicación: ').bold = True
            p.add_run(paciente['indicacion'])
        
        doc.add_page_break()
        
        # Código #Enzian
        doc.add_heading('CLASIFICACIÓN #ENZIAN', level=1)
        codigo = generar_codigo_enzian()
        p = doc.add_paragraph()
        p.add_run('Código: ').bold = True
        run = p.add_run(codigo)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 0, 128)
        
        doc.add_paragraph()
        
        # HALLAZGOS DETALLADOS
        doc.add_heading('HALLAZGOS DETALLADOS', level=1)
        
        # Peritoneo
        peritoneo = st.session_state.data['peritoneo']
        doc.add_heading('Peritoneo (P)', level=2)
        
        if peritoneo.get('estado') == 'anormal':
            p = doc.add_paragraph()
            p.add_run('Se identifican lesiones peritoneales superficiales. ')
            p.add_run(f"Clasificación: {peritoneo.get('clasificacion', 'N/A')}. ")
            
            if peritoneo.get('localizaciones'):
                p.add_run(f"Localizaciones: {', '.join(peritoneo['localizaciones'])}. ")
            
            if peritoneo.get('descripcion'):
                p.add_run(peritoneo['descripcion'])
        else:
            doc.add_paragraph('Sin evidencia de lesiones peritoneales superficiales.')
        
        # Ovarios
        doc.add_heading('Ovarios (O)', level=2)
        
        # Ovario derecho
        ovario_der = st.session_state.data['ovarios']['derecho']
        p = doc.add_paragraph()
        p.add_run('Ovario derecho: ').bold = True
        
        if ovario_der.get('estado') == 'anormal':
            p.add_run(f"Endometrioma de {ovario_der.get('diametro', 0)}cm. ")
            p.add_run(f"Clasificación: {ovario_der.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Estructura: {ovario_der.get('estructura', 'N/A')}. ")
            p.add_run(f"Contenido: {ovario_der.get('contenido', 'N/A')}. ")
            p.add_run(f"Vascularización: {ovario_der.get('vascularizacion', 'N/A')}. ")
            
            if ovario_der.get('adherencias'):
                p.add_run('Signos de adherencias a estructuras adyacentes. ')
                
            if ovario_der.get('descripcion'):
                p.add_run(ovario_der['descripcion'])
        elif ovario_der.get('estado') == 'no_visualizado':
            p.add_run('No visualizado.')
        else:
            p.add_run('Sin alteraciones evidentes.')
        
        # Ovario izquierdo
        ovario_izq = st.session_state.data['ovarios']['izquierdo']
        p = doc.add_paragraph()
        p.add_run('Ovario izquierdo: ').bold = True
        
        if ovario_izq.get('estado') == 'anormal':
            p.add_run(f"Endometrioma de {ovario_izq.get('diametro', 0)}cm. ")
            p.add_run(f"Clasificación: {ovario_izq.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Estructura: {ovario_izq.get('estructura', 'N/A')}. ")
            p.add_run(f"Contenido: {ovario_izq.get('contenido', 'N/A')}. ")
            p.add_run(f"Vascularización: {ovario_izq.get('vascularizacion', 'N/A')}. ")
            
            if ovario_izq.get('adherencias'):
                p.add_run('Signos de adherencias a estructuras adyacentes. ')
                
            if ovario_izq.get('descripcion'):
                p.add_run(ovario_izq['descripcion'])
        elif ovario_izq.get('estado') == 'no_visualizado':
            p.add_run('No visualizado.')
        else:
            p.add_run('Sin alteraciones evidentes.')
        
        # Condición tubo-ovárica
        doc.add_heading('Condición Tubo-Ovárica (T)', level=2)
        
        tubo_der = st.session_state.data['tubos']['derecho']
        p = doc.add_paragraph()
        p.add_run('Lado derecho: ').bold = True
        
        if tubo_der.get('estado') == 'anormal':
            p.add_run(f"{tubo_der.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Sliding sign: {tubo_der.get('sliding_sign', 'N/A')}. ")
            
            if tubo_der.get('permeabilidad') != 'No evaluada':
                p.add_run(f"Permeabilidad: {tubo_der.get('permeabilidad', 'N/A')}. ")
                
            if tubo_der.get('descripcion'):
                p.add_run(tubo_der['descripcion'])
        else:
            p.add_run('Movilidad preservada, sin adherencias evidentes.')
        
        tubo_izq = st.session_state.data['tubos']['izquierdo']
        p = doc.add_paragraph()
        p.add_run('Lado izquierdo: ').bold = True
        
        if tubo_izq.get('estado') == 'anormal':
            p.add_run(f"{tubo_izq.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Sliding sign: {tubo_izq.get('sliding_sign', 'N/A')}. ")
            
            if tubo_izq.get('permeabilidad') != 'No evaluada':
                p.add_run(f"Permeabilidad: {tubo_izq.get('permeabilidad', 'N/A')}. ")
                
            if tubo_izq.get('descripcion'):
                p.add_run(tubo_izq['descripcion'])
        else:
            p.add_run('Movilidad preservada, sin adherencias evidentes.')
        
        # Compartimento A
        doc.add_heading('Compartimento A (Vagina/Espacio Rectovaginal)', level=2)
        comp_a = st.session_state.data['compartimento_a']
        
        if comp_a.get('estado') == 'anormal':
            p = doc.add_paragraph()
            p.add_run(f"Lesión de endometriosis profunda de {comp_a.get('diametro', 0)}cm. ")
            p.add_run(f"Clasificación: {comp_a.get('clasificacion', 'N/A')}. ")
            
            if comp_a.get('localizacion'):
                p.add_run(f"Localización: {', '.join(comp_a['localizacion'])}. ")
                
            p.add_run(f"Ecogenicidad: {comp_a.get('ecogenicidad', 'N/A')}. ")
            p.add_run(f"Contornos: {comp_a.get('contornos', 'N/A')}. ")
            
            if comp_a.get('descripcion'):
                p.add_run(comp_a['descripcion'])
        else:
            doc.add_paragraph('Sin lesiones de endometriosis profunda en vagina ni espacio rectovaginal.')
        
        # Compartimento B
        doc.add_heading('Compartimento B (Ligamentos Uterosacros)', level=2)
        
        lsu_der = st.session_state.data['compartimento_b']['derecho']
        p = doc.add_paragraph()
        p.add_run('Ligamento uterosacro derecho: ').bold = True
        
        if lsu_der.get('estado') == 'anormal':
            p.add_run(f"Lesión de {lsu_der.get('diametro_max', 0)}cm ")
            p.add_run(f"(AP: {lsu_der.get('dim_ap', 0)}cm, CC: {lsu_der.get('dim_cc', 0)}cm). ")
            p.add_run(f"Clasificación: {lsu_der.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Sliding sign: {lsu_der.get('sliding_sign', 'N/A')}. ")
            
            if lsu_der.get('descripcion'):
                p.add_run(lsu_der['descripcion'])
        else:
            p.add_run('Sin alteraciones.')
        
        lsu_izq = st.session_state.data['compartimento_b']['izquierdo']
        p = doc.add_paragraph()
        p.add_run('Ligamento uterosacro izquierdo: ').bold = True
        
        if lsu_izq.get('estado') == 'anormal':
            p.add_run(f"Lesión de {lsu_izq.get('diametro_max', 0)}cm ")
            p.add_run(f"(AP: {lsu_izq.get('dim_ap', 0)}cm, CC: {lsu_izq.get('dim_cc', 0)}cm). ")
            p.add_run(f"Clasificación: {lsu_izq.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Sliding sign: {lsu_izq.get('sliding_sign', 'N/A')}. ")
            
            if lsu_izq.get('descripcion'):
                p.add_run(lsu_izq['descripcion'])
        else:
            p.add_run('Sin alteraciones.')
        
        # Compartimento C
        doc.add_heading('Compartimento C (Recto)', level=2)
        comp_c = st.session_state.data['compartimento_c']
        
        if comp_c.get('estado') == 'anormal':
            p = doc.add_paragraph()
            p.add_run(f"Lesión de endometriosis rectal de {comp_c.get('longitud', 0)}cm de longitud. ")
            p.add_run(f"Clasificación: {comp_c.get('clasificacion', 'N/A')}. ")
            p.add_run(f"Distancia desde margen anal: {comp_c.get('distancia_anal', 0)}cm. ")
            p.add_run(f"Profundidad de infiltración: {comp_c.get('profundidad', 'N/A')}. ")
            p.add_run(f"Circunferencia afectada: {comp_c.get('circunferencia', 0)}%. ")
            
            if comp_c.get('estenosis'):
                p.add_run('Signos de estenosis presentes. ')
                
            p.add_run(f"Sliding sign: {comp_c.get('sliding_sign', 'N/A')}. ")
            
            if comp_c.get('descripcion'):
                p.add_run(comp_c['descripcion'])
        else:
            doc.add_paragraph('Sin evidencia de endometriosis rectal.')
        
        # Localizaciones F
        doc.add_heading('Localizaciones Extragenitales (F)', level=2)
        loc_f = st.session_state.data['localizaciones_f']
        
        # Adenomiosis
        if loc_f.get('adenomiosis', {}).get('presente'):
            p = doc.add_paragraph()
            p.add_run('Adenomiosis (FA): ').bold = True
            
            criterios = loc_f['adenomiosis'].get('criterios_musa', [])
            if criterios:
                p.add_run(f"Criterios MUSA: {', '.join(criterios)}. ")
                
            if loc_f['adenomiosis'].get('descripcion'):
                p.add_run(loc_f['adenomiosis']['descripcion'])
        
        # Vejiga
        if loc_f.get('vejiga', {}).get('presente'):
            p = doc.add_paragraph()
            p.add_run('Vejiga (FB): ').bold = True
            
            vejiga_data = loc_f['vejiga']
            p.add_run(f"Lesión en {vejiga_data.get('localizacion', 'N/A')}. ")
            p.add_run(f"Profundidad: {vejiga_data.get('profundidad', 'N/A')}. ")
            p.add_run(f"Dimensión: {vejiga_data.get('dimension', 0)}cm. ")
            
            if vejiga_data.get('descripcion'):
                p.add_run(vejiga_data['descripcion'])
        
        # Uréter
        if loc_f.get('ureter', {}).get('presente'):
            p = doc.add_paragraph()
            p.add_run('Uréter (FU): ').bold = True
            
            ureter_data = loc_f['ureter']
            lados = ureter_data.get('lados', [])
            p.add_run(f"Compromiso ureteral {'bilateral' if len(lados) == 2 else lados[0].lower()}. ")
            p.add_run(f"Tipo: {ureter_data.get('tipo_compromiso', 'N/A')}. ")
            
            if ureter_data.get('descripcion'):
                p.add_run(ureter_data['descripcion'])
        
        # Intestino
        if loc_f.get('intestino', {}).get('presente'):
            p = doc.add_paragraph()
            p.add_run('Intestino (FI): ').bold = True
            
            intestino_data = loc_f['intestino']
            locs = intestino_data.get('localizaciones', [])
            p.add_run(f"Compromiso intestinal en: {', '.join(locs)}. ")
            p.add_run(f"Dimensión: {intestino_data.get('dimension', 0)}cm. ")
            
            if intestino_data.get('descripcion'):
                p.add_run(intestino_data['descripcion'])
        
        # Otras localizaciones
        if loc_f.get('otras', {}).get('presente'):
            p = doc.add_paragraph()
            p.add_run('Otras localizaciones: ').bold = True
            
            tipos = loc_f['otras'].get('tipos', [])
            p.add_run(f"{', '.join(tipos)}.")
        
        # Si no hay localizaciones F
        tiene_loc_f = any([
            loc_f.get('adenomiosis', {}).get('presente'),
            loc_f.get('vejiga', {}).get('presente'),
            loc_f.get('ureter', {}).get('presente'),
            loc_f.get('intestino', {}).get('presente'),
            loc_f.get('otras', {}).get('presente')
        ])
        
        if not tiene_loc_f:
            doc.add_paragraph('Sin compromiso de localizaciones extragenitales.')
        
        doc.add_page_break()
        
        # CONCLUSIONES
        doc.add_heading('CONCLUSIONES', level=1)
        
        p = doc.add_paragraph()
        p.add_run('Hallazgos ultrasonográficos compatibles con endometriosis según clasificación #Enzian:')
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(codigo)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 0, 128)
        
        doc.add_paragraph()
        
        # Recomendaciones
        doc.add_heading('RECOMENDACIONES', level=2)
        doc.add_paragraph('1. Correlación clínica con sintomatología de la paciente.')
        doc.add_paragraph('2. Valoración por especialista en endometriosis.')
        doc.add_paragraph('3. Considerar estudios complementarios según criterio clínico.')
        doc.add_paragraph('4. Planificación quirúrgica multidisciplinaria si está indicada.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Firma
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('_' * 50)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('Médico Ginecólogo')
        
        # Guardar en memoria
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    # Vista previa del reporte
    st.markdown("### 📊 Vista Previa del Código #Enzian")
    
    codigo_enzian = generar_codigo_enzian()
    
    st.markdown(f"""
    <div style="background-color: #f0f2f6; padding: 20px; border-radius: 10px; border-left: 5px solid #667eea;">
        <h3 style="color: #667eea;">Código Generado:</h3>
        <p style="font-size: 18px; color: #2c3e50; font-weight: bold;">{codigo_enzian}</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Validación de campos obligatorios
    st.markdown("### ✅ Validación de Datos")
    
    campos_obligatorios = []
    
    if not st.session_state.data['paciente'].get('nombre'):
        campos_obligatorios.append("❌ Nombre del paciente")
    else:
        st.success("✅ Nombre del paciente")
    
    if not st.session_state.data['paciente'].get('cedula'):
        campos_obligatorios.append("❌ Número de identificación")
    else:
        st.success("✅ Número de identificación")
    
    if not st.session_state.data['paciente'].get('fecha'):
        campos_obligatorios.append("❌ Fecha del estudio")
    else:
        st.success("✅ Fecha del estudio")
    
    # Verificar que al menos un compartimento tenga datos
    tiene_hallazgos = False
    
    if st.session_state.data['peritoneo'].get('estado') == 'anormal':
        tiene_hallazgos = True
    if any(ov.get('estado') == 'anormal' for ov in st.session_state.data['ovarios'].values()):
        tiene_hallazgos = True
    if any(tb.get('estado') == 'anormal' for tb in st.session_state.data['tubos'].values()):
        tiene_hallazgos = True
    if st.session_state.data['compartimento_a'].get('estado') == 'anormal':
        tiene_hallazgos = True
    if any(lsu.get('estado') == 'anormal' for lsu in st.session_state.data['compartimento_b'].values()):
        tiene_hallazgos = True
    if st.session_state.data['compartimento_c'].get('estado') == 'anormal':
        tiene_hallazgos = True
    if any(loc.get('presente') for loc in st.session_state.data['localizaciones_f'].values() if isinstance(loc, dict)):
        tiene_hallazgos = True
    
    if not tiene_hallazgos:
        st.warning("⚠️ No se han registrado hallazgos anormales en ningún compartimento")
    else:
        st.success("✅ Hallazgos registrados")
    
    st.markdown("---")
    
    # Botón para generar reporte
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        if campos_obligatorios:
            st.error("⚠️ Completa los campos obligatorios antes de generar el reporte:")
            for campo in campos_obligatorios:
                st.write(campo)
        else:
            if st.button("📄 GENERAR REPORTE EN WORD", type="primary", use_container_width=True):
                with st.spinner('Generando reporte...'):
                    buffer = generar_reporte_word()
                    
                    nombre_archivo = f"Reporte_Endometriosis_{st.session_state.data['paciente']['nombre'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                    
                    st.success("✅ ¡Reporte generado exitosamente!")
                    
                    st.download_button(
                        label="⬇️ DESCARGAR REPORTE",
                        data=buffer,
                        file_name=nombre_archivo,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
    
    st.markdown("---")
    
    # Resumen de hallazgos
    st.markdown("### 📋 Resumen de Hallazgos")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("#### Endometriosis Superficial")
        if st.session_state.data['peritoneo'].get('estado') == 'anormal':
            st.info(f"✓ Peritoneo: {st.session_state.data['peritoneo'].get('clasificacion', 'N/A')}")
        else:
            st.write("Sin hallazgos")
    
    with col2:
        st.markdown("#### Endometriosis Ovárica")
        ovario_izq = st.session_state.data['ovarios']['izquierdo']
        ovario_der = st.session_state.data['ovarios']['derecho']
        
        if ovario_izq.get('estado') == 'anormal':
            st.info(f"✓ Izquierdo: {ovario_izq.get('clasificacion', 'N/A')}")
        if ovario_der.get('estado') == 'anormal':
            st.info(f"✓ Derecho: {ovario_der.get('clasificacion', 'N/A')}")
        if ovario_izq.get('estado') != 'anormal' and ovario_der.get('estado') != 'anormal':
            st.write("Sin hallazgos")
    
    with col3:
        st.markdown("#### Adherencias")
        tubo_izq = st.session_state.data['tubos']['izquierdo']
        tubo_der = st.session_state.data['tubos']['derecho']
        
        if tubo_izq.get('estado') == 'anormal':
            clase = tubo_izq.get('clasificacion', 'N/A').split()[0]
            st.info(f"✓ Izquierdo: {clase}")
        if tubo_der.get('estado') == 'anormal':
            clase = tubo_der.get('clasificacion', 'N/A').split()[0]
            st.info(f"✓ Derecho: {clase}")
        if tubo_izq.get('estado') != 'anormal' and tubo_der.get('estado') != 'anormal':
            st.write("Sin hallazgos")
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("#### Endometriosis Profunda")
        hallazgos_de = []
        
        if st.session_state.data['compartimento_a'].get('estado') == 'anormal':
            hallazgos_de.append(f"✓ Compartimento A: {st.session_state.data['compartimento_a'].get('clasificacion', 'N/A')}")
        
        lsu_izq = st.session_state.data['compartimento_b']['izquierdo']
        lsu_der = st.session_state.data['compartimento_b']['derecho']
        
        if lsu_izq.get('estado') == 'anormal' or lsu_der.get('estado') == 'anormal':
            clase_izq = lsu_izq.get('clasificacion', 'B0')[1] if lsu_izq.get('estado') == 'anormal' else '0'
            clase_der = lsu_der.get('clasificacion', 'B0')[1] if lsu_der.get('estado') == 'anormal' else '0'
            hallazgos_de.append(f"✓ Compartimento B: B{clase_izq}/{clase_der}")
        
        if st.session_state.data['compartimento_c'].get('estado') == 'anormal':
            hallazgos_de.append(f"✓ Compartimento C: {st.session_state.data['compartimento_c'].get('clasificacion', 'N/A')}")
        
        if hallazgos_de:
            for hallazgo in hallazgos_de:
                st.info(hallazgo)
        else:
            st.write("Sin hallazgos")
    
    with col2:
        st.markdown("#### Localizaciones Asociadas")
        loc_f = st.session_state.data['localizaciones_f']
        hallazgos_f = []
        
        if loc_f.get('adenomiosis', {}).get('presente'):
            hallazgos_f.append("✓ Adenomiosis (FA)")
        if loc_f.get('vejiga', {}).get('presente'):
            hallazgos_f.append("✓ Vejiga (FB)")
        if loc_f.get('ureter', {}).get('presente'):
            lados = loc_f['ureter'].get('lados', [])
            hallazgos_f.append(f"✓ Uréter: {', '.join(lados)}")
        if loc_f.get('intestino', {}).get('presente'):
            hallazgos_f.append("✓ Intestino (FI)")
        if loc_f.get('otras', {}).get('presente'):
            hallazgos_f.append("✓ Otras localizaciones")
        
        if hallazgos_f:
            for hallazgo in hallazgos_f:
                st.info(hallazgo)
        else:
            st.write("Sin hallazgos")
    
    with col3:
        st.markdown("#### Alertas Clínicas")
        alertas = []
        
        # Alerta por hidronefrosis
        if loc_f.get('ureter', {}).get('presente'):
            alertas.append("⚠️ Compromiso ureteral - Valorar función renal")
        
        # Alerta por estenosis rectal
        if st.session_state.data['compartimento_c'].get('estenosis'):
            alertas.append("⚠️ Estenosis rectal presente")
        
        # Alerta por endometriomas grandes
        if ovario_izq.get('estado') == 'anormal' and ovario_izq.get('diametro', 0) > 7:
            alertas.append("⚠️ Endometrioma izquierdo >7cm")
        if ovario_der.get('estado') == 'anormal' and ovario_der.get('diametro', 0) > 7:
            alertas.append("⚠️ Endometrioma derecho >7cm")
        
        # Alerta por enfermedad extensa
        if st.session_state.data['compartimento_a'].get('clasificacion') in ['A3 (>3 cm)', 'A3']:
            alertas.append("⚠️ Endometriosis profunda extensa (A3)")
        
        if alertas:
            for alerta in alertas:
                st.warning(alerta)
        else:
            st.success("✅ Sin alertas críticas")

# Footer
st.markdown("---")
st.markdown("""
    <div style="text-align: center; color: #666; padding: 20px;">
        <p><strong>Sistema de Reporte Ultrasonográfico de Endometriosis</strong></p>
        <p>Clasificación #Enzian - Versión 2021</p>
        <p style="font-size: 0.9em;">Desarrollado para evaluación sistemática de endometriosis mediante ultrasonido transvaginal</p>
    </div>
""", unsafe_allow_html=True)
